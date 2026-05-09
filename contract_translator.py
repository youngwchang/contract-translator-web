#!/usr/bin/env python3
"""
영문 계약서 번역 자동화 스크립트
- 로컬 PDF / DOCX 파일 읽기 (이미지 PDF는 Claude Vision OCR 자동 처리)
- Track Changes 자동 처리 (최종본 기준으로 추출)
- Anthropic Claude API로 조항 추출 및 한국어 번역
- 결과를 Excel (.xlsx) 로 저장

사용법:
    python contract_translator.py                     # 대화형 실행
    python contract_translator.py contract.pdf        # 파일 직접 지정
    python contract_translator.py contract.docx -o output.xlsx
"""

import sys
import os
import json
import argparse
import re
import base64
from pathlib import Path
from datetime import datetime


# ── 의존성 체크 ──────────────────────────────────────────────────────────────

def check_and_import():
    missing = []
    try:
        import anthropic
    except ImportError:
        missing.append("anthropic")
    try:
        import pdfplumber
    except ImportError:
        missing.append("pdfplumber")
    try:
        import docx
    except ImportError:
        missing.append("python-docx")
    try:
        import openpyxl
    except ImportError:
        missing.append("openpyxl")

    if missing:
        print(f"\n[오류] 필수 패키지가 설치되지 않았습니다: {', '.join(missing)}")
        print("아래 명령어로 설치하세요:\n")
        print(f"    pip install {' '.join(missing)}\n")
        sys.exit(1)

check_and_import()

import anthropic
import pdfplumber
import docx
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


# ── Track Changes 처리 ────────────────────────────────────────────────────────

def accept_tracked_changes(doc) -> None:
    """Track Changes 해소: 삽입(ins)은 유지, 삭제(del)는 제거"""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body

    # 삭제된 내용 제거 (w:del 요소 전체 제거)
    for el in list(body.iter(f'{{{W}}}del')):
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    # 삽입 래퍼 제거 (w:ins 요소는 제거하고 자식은 유지)
    for el in reversed(list(body.iter(f'{{{W}}}ins'))):
        parent = el.getparent()
        if parent is None:
            continue
        idx = list(parent).index(el)
        for child in list(el):
            parent.insert(idx, child)
            idx += 1
        parent.remove(el)


# ── 텍스트 추출 ───────────────────────────────────────────────────────────────

def extract_pdf(path: Path) -> str:
    """PDF에서 텍스트 추출 (텍스트 레이어 있는 PDF 전용)"""
    print(f"  → PDF 텍스트 추출 중: {path.name}")
    pages = []
    with pdfplumber.open(path) as pdf:
        total = len(pdf.pages)
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text:
                pages.append(text.strip())
            if i % 10 == 0:
                print(f"     {i}/{total} 페이지 처리 완료...")
    full_text = "\n\n".join(pages)
    print(f"  → 추출 완료: {len(pages)} 페이지, {len(full_text):,} 자")
    return full_text


def extract_pdf_ocr(path: Path, client: anthropic.Anthropic) -> str:
    """이미지 PDF에서 Claude Vision으로 OCR 텍스트 추출"""
    try:
        import fitz
    except ImportError:
        print("\n[오류] PyMuPDF가 설치되지 않았습니다.")
        print("아래 명령어로 설치하세요:\n  pip install PyMuPDF\n")
        sys.exit(1)

    print(f"  → 이미지(스캔) PDF 감지: Claude Vision OCR 시작")
    doc = fitz.open(str(path))
    total_pages = len(doc)
    print(f"  → 총 {total_pages}페이지 OCR 처리 예정")

    extracted_pages = []
    BATCH = 4

    for start in range(0, total_pages, BATCH):
        end = min(start + BATCH, total_pages)
        print(f"     OCR {start+1}~{end}/{total_pages} 페이지...")

        content = []
        for pn in range(start, end):
            page = doc[pn]
            pix = page.get_pixmap(dpi=150)
            img_b64 = base64.standard_b64encode(pix.tobytes("png")).decode()
            content.append({"type": "text", "text": f"--- Page {pn + 1} ---"})
            content.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": "image/png",
                    "data": img_b64
                }
            })

        content.append({
            "type": "text",
            "text": (
                "Extract ALL text from the above contract page images verbatim. "
                "Preserve line breaks, indentation, numbering, and document structure exactly. "
                "Keep the '--- Page N ---' separators between pages. "
                "Return only the extracted text, nothing else."
            )
        })

        msg = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=8192,
            messages=[{"role": "user", "content": content}]
        )
        extracted_pages.append(msg.content[0].text.strip())

    doc.close()
    full_text = "\n\n".join(extracted_pages)
    print(f"  → OCR 완료: {total_pages}페이지, {len(full_text):,}자 추출")
    return full_text


_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _cell_text(tc_elem) -> str:
    """<w:tc> 요소에서 텍스트 추출 — 중첩 테이블 포함, 재귀 처리"""
    parts = []
    for child in tc_elem:
        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if local == 'p':
            t = ''.join(n.text or '' for n in child.iter(f'{{{_W}}}t')).strip()
            if t:
                parts.append(t)
        elif local == 'tbl':
            parts.extend(_table_rows(child))
    return '\n'.join(parts)


def _table_rows(tbl_elem) -> list[str]:
    """<w:tbl> 요소를 텍스트 행 리스트로 변환 — 중첩 테이블 재귀 처리"""
    rows = []
    for tr in tbl_elem.findall(f'{{{_W}}}tr'):
        cells = []
        for tc in tr.findall(f'{{{_W}}}tc'):
            text = _cell_text(tc)
            if text:
                cells.append(text)
        if cells:
            rows.append(' | '.join(cells))
    return rows


def extract_docx(path: Path) -> str:
    """DOCX에서 텍스트 추출 (Track Changes 처리 + 중첩 테이블 포함)"""
    print(f"  → DOCX 텍스트 추출 중: {path.name}")
    doc = docx.Document(path)

    # Track Changes 처리
    try:
        accept_tracked_changes(doc)
        print("  → Track Changes 처리: 최종본 기준으로 추출")
    except Exception as e:
        print(f"  → Track Changes 처리 건너뜀: {e}")

    # 문서 본문을 순서대로 순회 (단락 + 최상위 테이블 모두 포함)
    parts = []
    for child in doc.element.body:
        local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if local == 'p':
            text = ''.join(n.text or '' for n in child.iter(f'{{{_W}}}t')).strip()
            if text:
                parts.append(text)
        elif local == 'tbl':
            rows = _table_rows(child)
            parts.extend(rows)

    full_text = '\n\n'.join(parts)
    print(f"  → 추출 완료: {len(parts)} 항목, {len(full_text):,} 자")
    return full_text


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path)
    elif ext in (".docx", ".doc"):
        return extract_docx(path)
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {ext}  (PDF, DOCX만 가능)")


# ── Claude API 호출 ───────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are a professional legal contract analyst and Korean translator specializing in pharmaceutical licensing and distribution agreements.

Extract every article, clause, section, and provision from the contract. For each one provide:
1. clause: The identifier (e.g. "Article 1", "Section 2.3", "Schedule A", "Recital", "Definitions")
2. original: The complete original English text of that clause (verbatim, do not summarize)
3. korean: A precise, professional Korean translation

Return ONLY a valid JSON array. No markdown fences, no preamble, no explanation.
Format exactly:
[{"clause":"Article 1","original":"Full English text...","korean":"정확한 한국어 번역..."}]

Rules:
- Include ALL parts: preamble, recitals, definitions, main articles, schedules, appendices, signature blocks
- Do not merge or split clauses — extract as written
- Keep subsections (1.1, 1.2...) as separate entries if they have distinct content
- Translate with legal precision; use standard Korean legal terminology
- For tables: include table data as part of the relevant clause text

Before finalizing your response, double-check all Korean legal terminology for accuracy and consistency. Verify that pharmaceutical/licensing-specific terms use standard Korean legal expressions (e.g., 라이선스, 독점적 실시권, 계약 해지, 손해배상 등)."""


def recover_partial_json(raw: str) -> list[dict]:
    """잘린 JSON 응답에서 완성된 객체만 추출"""
    results = []
    decoder = json.JSONDecoder()
    i = 0
    while i < len(raw):
        if raw[i] == '{':
            try:
                obj, next_i = decoder.raw_decode(raw, i)
                if isinstance(obj, dict) and ('clause' in obj or 'category' in obj):
                    results.append(obj)
                i = next_i
            except json.JSONDecodeError:
                i += 1
        else:
            i += 1
    return results


GLOSSARY_SYSTEM_PROMPT = """You are a legal terminology specialist for pharmaceutical licensing and distribution contracts.

Extract a glossary of recurring key terms from the contract. For each term provide:
- "english": exact English term as it appears (e.g., "Minimum Purchase Quantity")
- "korean": the standard Korean translation to use consistently throughout
- "category": one of "당사자" | "정의/제품" | "권리/의무" | "재무" | "기간/해지" | "기타"

Rules:
- Include only terms that (a) appear MULTIPLE TIMES, or (b) are capitalized defined terms (e.g., "the Product", "Net Sales"), or (c) are critical legal/business concepts
- Use standard Korean legal and pharmaceutical terminology
- Maximum 50 most important terms — prioritize defined terms and recurring legal phrases
- Examples to look for:
  • Parties: Licensor, Licensee, Distributor, Manufacturer, Supplier
  • Product/territory: Product, Territory, Field, Indication
  • Financial: Net Sales, Royalty, Milestone Payment, Minimum Purchase Quantity, Transfer Price
  • Time/term: Initial Term, Renewal Term, Notice Period, Effective Date
  • IP: Patent Rights, Know-How, Trademark, Confidential Information
  • Termination: Material Breach, Change of Control, Cure Period

Return ONLY a valid JSON array. No markdown fences, no preamble.
Format: [{"english":"Minimum Purchase Quantity","korean":"최소구매수량","category":"재무"}]"""


def extract_glossary(client: anthropic.Anthropic, text: str) -> list[dict]:
    """청크 분할 번역 시 용어 일관성을 위해 핵심 용어 사전을 추출"""
    print("  → 핵심 용어 사전 추출 중 (청크 간 일관된 번역을 위해)...")

    # 정의는 앞쪽, 핵심 의무는 뒤쪽에 주로 등장 → 양쪽을 샘플링
    MAX_CHARS = 80000
    if len(text) > MAX_CHARS:
        half = MAX_CHARS // 2
        sample = text[:half] + "\n\n[... 중략 ...]\n\n" + text[-half:]
    else:
        sample = text

    try:
        message = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=4096,
            system=GLOSSARY_SYSTEM_PROMPT,
            messages=[{
                "role": "user",
                "content": f"Contract text:\n\n{sample}\n\nExtract glossary as a JSON array."
            }]
        )
        raw = message.content[0].text.strip()
        raw = re.sub(r'^```json\s*', '', raw)
        raw = re.sub(r'^```\s*', '', raw)
        raw = re.sub(r'\s*```$', '', raw)
        glossary = json.loads(raw.strip())
        if not isinstance(glossary, list):
            raise ValueError("응답이 리스트 형식이 아닙니다")
        print(f"  → {len(glossary)}개 핵심 용어 추출 완료")
        # 미리보기 (상위 8개)
        for term in glossary[:8]:
            en = term.get("english", "")
            ko = term.get("korean", "")
            if en and ko:
                print(f"     · {en} → {ko}")
        if len(glossary) > 8:
            print(f"     · ... 외 {len(glossary) - 8}개")
        return glossary
    except Exception as e:
        print(f"  [경고] 용어 사전 추출 실패: {e}. 일반 번역으로 진행합니다.")
        return []


def format_glossary_for_prompt(glossary: list[dict]) -> str:
    """용어 사전을 system prompt에 주입할 형식으로 변환"""
    if not glossary:
        return ""
    lines = [
        "",
        "=== MANDATORY TERMINOLOGY (use these Korean translations consistently) ===",
        "When you encounter any of these English terms, translate them using EXACTLY the Korean equivalent below. Do not paraphrase or vary the translation.",
        "",
    ]
    for term in glossary:
        en = (term.get("english") or "").strip()
        ko = (term.get("korean") or "").strip()
        if en and ko:
            lines.append(f'  • "{en}" → "{ko}"')
    lines.append("=== END TERMINOLOGY ===")
    return "\n".join(lines)


def chunk_text(text: str, max_chars: int = 15000) -> list[str]:
    """긴 계약서를 청크로 분할 (조항 경계 기준, 기본 15000자)"""
    if len(text) <= max_chars:
        return [text]

    chunks = []
    current = ""
    splitter = re.compile(r'(?=\b(?:Article|Section|ARTICLE|SECTION|Clause|CLAUSE)\s+\d)', re.IGNORECASE)
    parts = splitter.split(text)

    for part in parts:
        if len(current) + len(part) > max_chars and current:
            chunks.append(current.strip())
            current = part
        else:
            current += "\n" + part

    if current.strip():
        chunks.append(current.strip())

    print(f"  → 계약서가 길어 {len(chunks)}개 청크로 분할 처리합니다")
    return chunks


def call_api(client: anthropic.Anthropic, text: str, chunk_index: int = 0, total_chunks: int = 1, glossary: list[dict] | None = None) -> list[dict]:
    """Claude API 호출 → 조항 리스트 반환 (잘린 응답 복구 포함)"""
    suffix = f" (청크 {chunk_index+1}/{total_chunks})" if total_chunks > 1 else ""
    print(f"  → Claude API 호출 중{suffix}...")

    user_content = f"Contract text:\n\n{text}\n\nExtract all clauses and return as a JSON array."
    if total_chunks > 1:
        user_content = (
            f"[This is part {chunk_index+1} of {total_chunks} of a large contract.]\n\n"
            + user_content
        )

    system_prompt = SYSTEM_PROMPT + format_glossary_for_prompt(glossary or [])

    message = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=8192,
        system=system_prompt,
        messages=[{"role": "user", "content": user_content}]
    )

    raw = message.content[0].text.strip()
    stop_reason = message.stop_reason

    # 마크다운 펜스 제거
    raw = re.sub(r'^```json\s*', '', raw)
    raw = re.sub(r'^```\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    raw = raw.strip()

    # 정상 파싱 시도
    try:
        result = json.loads(raw)
        print(f"  → {len(result)}개 조항 추출 완료{suffix}")
        return result
    except json.JSONDecodeError:
        pass

    # 응답이 max_tokens로 잘린 경우 부분 복구
    if stop_reason == "max_tokens":
        print(f"  [경고] 응답이 토큰 한도로 잘렸습니다{suffix}. 완성된 조항만 저장합니다.")
        recovered = recover_partial_json(raw)
        print(f"  → 부분 복구: {len(recovered)}개 조항{suffix}")
        return recovered

    # 기타 파싱 실패
    print(f"  [경고] JSON 파싱 실패{suffix}. 부분 복구를 시도합니다.")
    recovered = recover_partial_json(raw)
    if recovered:
        print(f"  → 부분 복구: {len(recovered)}개 조항{suffix}")
        return recovered

    print(f"  응답 미리보기: {raw[:200]}...")
    return []


def translate_contract(client: anthropic.Anthropic, text: str) -> list[dict]:
    """전체 계약서 번역 처리 (청크 분할 + 용어 일관성 글로서리)"""
    chunks = chunk_text(text)

    # 청크가 2개 이상일 때만 글로서리를 추출 (단일 청크는 자체적으로 일관됨)
    glossary = []
    if len(chunks) > 1:
        glossary = extract_glossary(client, text)

    all_results = []
    for i, chunk in enumerate(chunks):
        results = call_api(client, chunk, i, len(chunks), glossary=glossary)
        all_results.extend(results)
    return all_results


# ── Excel 저장 ────────────────────────────────────────────────────────────────

COLOR_HEADER_BG = "1A3A6C"
COLOR_HEADER_FG = "FFFFFF"
COLOR_ROW_ODD   = "FFFFFF"
COLOR_ROW_EVEN  = "F7F9FC"
COLOR_BORDER    = "D0D7E3"
COLOR_CLAUSE_BG = "DBEAFE"
COLOR_CLAUSE_FG = "1E40AF"


def thin_border():
    side = Side(style="thin", color=COLOR_BORDER)
    return Border(left=side, right=side, top=side, bottom=side)


def save_excel(results: list[dict], output_path: Path, source_filename: str):
    # 파일이 이미 열려 있으면 번호를 붙인 새 이름으로 저장
    actual_path = output_path
    counter = 1
    while actual_path.exists():
        try:
            actual_path.open("r+b").close()
            break  # 열 수 있으면 잠금 없음 → 덮어써도 됨
        except PermissionError:
            actual_path = output_path.parent / f"{output_path.stem}_{counter}{output_path.suffix}"
            counter += 1
    output_path = actual_path

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "계약서 번역"

    ws.merge_cells("A1:C1")
    title_cell = ws["A1"]
    title_cell.value = f"영문 계약서 번역본 | {source_filename} | {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    title_cell.font = Font(name="맑은 고딕", size=10, color="6B7280", italic=True)
    title_cell.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 20

    headers = ["조항 (Clause)", "원문 (Original English)", "한국어 번역 (Korean Translation)"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col, value=header)
        cell.font = Font(name="맑은 고딕", size=10, bold=True, color=COLOR_HEADER_FG)
        cell.fill = PatternFill("solid", fgColor=COLOR_HEADER_BG)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border()
    ws.row_dimensions[2].height = 28

    for i, item in enumerate(results):
        row = i + 3
        is_even = i % 2 == 1
        row_bg = COLOR_ROW_EVEN if is_even else COLOR_ROW_ODD

        clause   = item.get("clause", "")
        original = item.get("original", "")
        korean   = item.get("korean", "")

        c1 = ws.cell(row=row, column=1, value=clause)
        c1.font = Font(name="맑은 고딕", size=9, bold=True, color=COLOR_CLAUSE_FG)
        c1.fill = PatternFill("solid", fgColor=COLOR_CLAUSE_BG if not is_even else "E0EDFF")
        c1.alignment = Alignment(horizontal="center", vertical="top", wrap_text=True)
        c1.border = thin_border()

        c2 = ws.cell(row=row, column=2, value=original)
        c2.font = Font(name="Calibri", size=9, color="334155")
        c2.fill = PatternFill("solid", fgColor=row_bg)
        c2.alignment = Alignment(vertical="top", wrap_text=True)
        c2.border = thin_border()

        c3 = ws.cell(row=row, column=3, value=korean)
        c3.font = Font(name="맑은 고딕", size=9, color="0F1F3D")
        c3.fill = PatternFill("solid", fgColor=row_bg)
        c3.alignment = Alignment(vertical="top", wrap_text=True)
        c3.border = thin_border()

        max_len = max(len(original), len(korean))
        ws.row_dimensions[row].height = min(max(40, max_len // 6), 200)

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 65
    ws.column_dimensions["C"].width = 65
    ws.freeze_panes = "A3"
    ws.auto_filter.ref = f"A2:C{len(results) + 2}"
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.print_title_rows = "2:2"

    try:
        wb.save(output_path)
    except PermissionError:
        # 저장 직전에 파일이 열린 경우 대비
        stem = output_path.stem.rstrip("_0123456789")
        for i in range(1, 20):
            alt = output_path.parent / f"{stem}_{i}{output_path.suffix}"
            if not alt.exists():
                wb.save(alt)
                output_path = alt
                break
        else:
            raise
    print(f"\n  ✓ Excel 저장 완료: {output_path}")
    print(f"    - 총 {len(results)}개 조항")
    print(f"    - 파일 크기: {output_path.stat().st_size / 1024:.1f} KB")


# ── API 키 관리 ───────────────────────────────────────────────────────────────

def get_api_key() -> str:
    key = os.environ.get("ANTHROPIC_API_KEY", "")
    if key:
        print("  → API 키: 환경변수(ANTHROPIC_API_KEY)에서 로드")
        return key

    env_file = Path(__file__).parent / ".env"
    if not env_file.exists():
        env_file = Path(".env")
    if env_file.exists():
        for line in env_file.read_text(encoding="utf-8").splitlines():
            if line.startswith("ANTHROPIC_API_KEY="):
                key = line.split("=", 1)[1].strip().strip('"').strip("'")
                if key:
                    print("  → API 키: .env 파일에서 로드")
                    return key

    print("\n  Anthropic API 키를 입력하세요 (입력값은 화면에 표시되지 않습니다):")
    import getpass
    key = getpass.getpass("  ANTHROPIC_API_KEY: ").strip()
    if not key:
        print("[오류] API 키가 입력되지 않았습니다.")
        sys.exit(1)
    return key


# ── 핵심조항 분석 (계약기간·해지 관련) ──────────────────────────────────────

KEY_TERMS_SYSTEM_PROMPT = """You are a legal analyst for a Korean pharmaceutical company reviewing licensing/distribution agreements.

Extract ONLY these critical contract terms and return a JSON array.

Each item must have:
- "category": one of exactly these values:
    "계약기간" | "갱신·연장 조건" | "해지 통보" | "해지 사유" | "기타 핵심 조항"
- "item": specific term name in Korean (e.g., "최소구매수량 미달성", "서면 통지 기간")
- "clause_ref": article/section number (e.g., "Article 12.3")
- "original": the exact English text from the contract (verbatim)
- "korean": clear, precise Korean explanation

Categories to extract:
1. 계약기간 — contract start/end dates, duration
2. 갱신·연장 조건 — how/when contract renews or extends, auto-renewal clauses
3. 해지 통보 — notice requirements (written/email, notice period in days/months, delivery method)
4. 해지 사유 — ALL grounds for termination:
   • Minimum purchase quantity failure (최소구매수량 미달성)
   • Material breach (중대한 계약 위반)
   • Insolvency / bankruptcy (파산·지급불능)
   • Regulatory approval issues (허가 취소·거부)
   • Change of control
   • Any other termination trigger
5. 기타 핵심 조항 — milestone payments, penalties, cure periods

Before finalizing your response, double-check all Korean legal terminology for accuracy and consistency. Verify that pharmaceutical/licensing-specific terms use standard Korean legal expressions (e.g., 계약기간, 해지 통보, 최소구매수량, 중대한 위반 등).

Return ONLY a valid JSON array. No markdown, no preamble."""


def analyze_key_terms(client: anthropic.Anthropic, text: str) -> list[dict]:
    """계약 핵심 조항 분석 (계약기간·해지 관련) — 단일 API 호출"""
    print("  → 핵심 조항 추출 중 (계약기간, 갱신, 해지)...")

    # 전체 문서를 한 번에 전송 (출력이 짧으므로 입력이 길어도 무방)
    MAX_CHARS = 120000
    if len(text) > MAX_CHARS:
        print(f"  → 문서가 길어 앞부분 {MAX_CHARS//1000}K자 기준으로 분석합니다")
        text = text[:MAX_CHARS]

    message = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=8192,
        system=KEY_TERMS_SYSTEM_PROMPT,
        messages=[{
            "role": "user",
            "content": f"Contract text:\n\n{text}\n\nExtract all key terms as a JSON array."
        }]
    )

    raw = message.content[0].text.strip()
    raw = re.sub(r'^```json\s*', '', raw)
    raw = re.sub(r'^```\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)

    try:
        result = json.loads(raw.strip())
        print(f"  → {len(result)}개 핵심 조항 추출 완료")
        return result
    except json.JSONDecodeError:
        recovered = recover_partial_json(raw)
        if recovered:
            print(f"  → 부분 복구: {len(recovered)}개 조항")
            return recovered
        print(f"  [경고] 파싱 실패. 응답 미리보기: {raw[:300]}")
        return []


# 카테고리별 색상
_CATEGORY_COLORS = {
    "계약기간":     ("1E3A8A", "DBEAFE"),  # 파랑
    "갱신·연장 조건": ("14532D", "DCFCE7"),  # 초록
    "해지 통보":    ("7C2D12", "FED7AA"),  # 주황
    "해지 사유":    ("7F1D1D", "FEE2E2"),  # 빨강
    "기타 핵심 조항": ("1F2937", "F3F4F6"),  # 회색
}
_DEFAULT_COLOR = ("374151", "F9FAFB")


def save_key_terms_excel(results: list[dict], output_path: Path, source_filename: str):
    """핵심 조항 분석 결과를 Excel로 저장"""
    # 파일 잠금 처리
    actual_path = output_path
    counter = 1
    while actual_path.exists():
        try:
            actual_path.open("r+b").close()
            break
        except PermissionError:
            actual_path = output_path.parent / f"{output_path.stem}_{counter}{output_path.suffix}"
            counter += 1
    output_path = actual_path

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "핵심 조항 요약"

    # 제목 행
    ws.merge_cells("A1:E1")
    t = ws["A1"]
    t.value = f"계약 핵심 조항 요약 | {source_filename} | {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    t.font = Font(name="맑은 고딕", size=10, italic=True, color="6B7280")
    t.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 20

    # 헤더 행
    headers = ["분류", "항목", "조항 번호", "원문 (English)", "한국어 설명"]
    col_widths = [18, 28, 14, 55, 55]
    for col, (header, width) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=2, column=col, value=header)
        cell.font = Font(name="맑은 고딕", size=10, bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1A3A6C")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border()
        ws.column_dimensions[get_column_letter(col)].width = width
    ws.row_dimensions[2].height = 28

    # 데이터 행
    for i, item in enumerate(results):
        row = i + 3
        category = item.get("category", "기타 핵심 조항")
        fg_color, bg_color = _CATEGORY_COLORS.get(category, _DEFAULT_COLOR)

        values = [
            category,
            item.get("item", ""),
            item.get("clause_ref", ""),
            item.get("original", ""),
            item.get("korean", ""),
        ]
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.fill = PatternFill("solid", fgColor=bg_color)
            cell.border = thin_border()
            cell.alignment = Alignment(vertical="top", wrap_text=True,
                                       horizontal="center" if col <= 3 else "left")
            if col == 1:  # 분류 열: 굵게 + 카테고리 색
                cell.font = Font(name="맑은 고딕", size=9, bold=True, color=fg_color)
            elif col == 2:
                cell.font = Font(name="맑은 고딕", size=9, bold=True, color="0F1F3D")
            elif col == 3:
                cell.font = Font(name="Calibri", size=9, bold=True, color=fg_color)
            elif col == 4:
                cell.font = Font(name="Calibri", size=9, color="334155")
            else:
                cell.font = Font(name="맑은 고딕", size=9, color="0F1F3D")

        max_len = max(len(item.get("original", "")), len(item.get("korean", "")))
        ws.row_dimensions[row].height = min(max(30, max_len // 8), 150)

    ws.freeze_panes = "A3"
    ws.auto_filter.ref = f"A2:E{len(results) + 2}"
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.print_title_rows = "2:2"

    try:
        wb.save(output_path)
    except PermissionError:
        stem = output_path.stem.rstrip("_0123456789")
        for i in range(1, 20):
            alt = output_path.parent / f"{stem}_{i}{output_path.suffix}"
            if not alt.exists():
                wb.save(alt)
                output_path = alt
                break

    print(f"\n  ✓ 핵심 조항 요약 저장 완료: {output_path}")
    print(f"    - 총 {len(results)}개 항목")

    # 카테고리별 요약 출력
    from collections import Counter
    counts = Counter(item.get("category", "") for item in results)
    for cat, cnt in sorted(counts.items()):
        print(f"    - {cat}: {cnt}개")


# ── Word 저장 ─────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """python-docx 셀 배경색 설정 (XML 직접 조작)"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "D0D7E3"):
    """셀 테두리 설정"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _word_cell(cell, text: str, bold: bool = False,
               font_name: str = "맑은 고딕", font_size_pt: int = 9,
               color_hex: str = "000000", align: str = "left",
               bg_hex: str | None = None, border_color: str = "D0D7E3"):
    """셀에 텍스트 + 서식 적용 헬퍼"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    align_map = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)

    run = para.add_run(str(text) if text else "")
    run.bold      = bold
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)

    if bg_hex:
        _set_cell_bg(cell, bg_hex)
    _set_cell_borders(cell, border_color)


def _set_landscape(doc):
    """문서를 가로 방향으로 설정 (A4 landscape)"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Mm
    section = doc.sections[0]
    # A4 landscape: 297mm × 210mm
    section.page_width  = Mm(297)
    section.page_height = Mm(210)
    section.left_margin   = Mm(15)
    section.right_margin  = Mm(15)
    section.top_margin    = Mm(15)
    section.bottom_margin = Mm(15)
    # XML 상 orientation 속성
    pgSz = section._sectPr.find(qn("w:pgSz"))
    if pgSz is not None:
        pgSz.set(qn("w:orient"), "landscape")


def save_word(results: list[dict], output_path: Path, source_filename: str):
    """전체 번역 결과를 Word(.docx) 로 저장
    컬럼: 조항 | 원문(English) | 한국어 번역
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _set_landscape(doc)

    # ── 제목 단락 ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(
        f"영문 계약서 전체 번역  |  {source_filename}  |  {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.name  = "맑은 고딕"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.italic = True
    title.paragraph_format.space_after = Pt(6)

    # ── 테이블 생성 ──
    # A4 landscape 가용 폭(mm): 297 - 15 - 15 = 267mm → DXA: 267*56.7 ≈ 15138
    # 컬럼 비율 — 조항:원문:한국어 = 18:41:41
    TOTAL_DXA = 15100
    col_w = [int(TOTAL_DXA * r) for r in (0.18, 0.41, 0.41)]

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.autofit = False
    from docx.shared import Twips
    tbl.width = Twips(TOTAL_DXA)
    for i, w in enumerate(col_w):
        tbl.columns[i].width = Twips(w)

    # 헤더 행
    hdr_row  = tbl.rows[0]
    hdr_row.height = Twips(450)
    headers  = ["조항 (Clause)", "원문 (Original English)", "한국어 번역 (Korean Translation)"]
    for ci, (cell, hdr) in enumerate(zip(hdr_row.cells, headers)):
        cell.width = Twips(col_w[ci])
        _word_cell(cell, hdr, bold=True, font_name="맑은 고딕", font_size_pt=9,
                   color_hex="FFFFFF", align="center",
                   bg_hex="1A3A6C", border_color="1A3A6C")

    # 데이터 행
    for i, item in enumerate(results):
        clause   = item.get("clause", "")
        original = item.get("original", "")
        korean   = item.get("korean", "")
        bg = "F7F9FC" if i % 2 == 1 else "FFFFFF"

        row = tbl.add_row()
        row.cells[0].width = Twips(col_w[0])
        row.cells[1].width = Twips(col_w[1])
        row.cells[2].width = Twips(col_w[2])

        _word_cell(row.cells[0], clause, bold=True, font_size_pt=9,
                   color_hex="1E40AF", align="center", bg_hex="DBEAFE" if i % 2 == 0 else "E0EDFF")
        _word_cell(row.cells[1], original, font_name="Calibri", font_size_pt=9,
                   color_hex="334155", bg_hex=bg)
        _word_cell(row.cells[2], korean, font_size_pt=9,
                   color_hex="0F1F3D", bg_hex=bg)

    doc.save(str(output_path))
    print(f"\n  ✓ Word 저장 완료: {output_path}")
    print(f"    - 총 {len(results)}개 조항")


_KT_COLORS_WORD = {
    "계약기간":       ("1E3A8A", "DBEAFE"),
    "갱신·연장 조건": ("14532D", "DCFCE7"),
    "해지 통보":      ("7C2D12", "FED7AA"),
    "해지 사유":      ("7F1D1D", "FEE2E2"),
    "기타 핵심 조항": ("1F2937", "F3F4F6"),
}
_KT_DEFAULT_WORD = ("374151", "F9FAFB")


def save_key_terms_word(results: list[dict], output_path: Path, source_filename: str):
    """핵심 조항 분석 결과를 Word(.docx) 로 저장
    컬럼: 분류 | 항목 | 조항번호 | 원문(English) | 한국어 설명
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _set_landscape(doc)

    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(
        f"계약 핵심 조항 요약  |  {source_filename}  |  {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.name  = "맑은 고딕"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.italic = True
    title.paragraph_format.space_after = Pt(6)

    TOTAL_DXA = 15100
    # 분류:항목:조항번호:원문:한국어 = 14:20:10:28:28
    col_w = [int(TOTAL_DXA * r) for r in (0.14, 0.20, 0.10, 0.28, 0.28)]

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    tbl.autofit = False
    tbl.width = Twips(TOTAL_DXA)
    for i, w in enumerate(col_w):
        tbl.columns[i].width = Twips(w)

    hdr_row = tbl.rows[0]
    hdr_row.height = Twips(450)
    headers = ["분류", "항목", "조항 번호", "원문 (English)", "한국어 설명"]
    for ci, (cell, hdr) in enumerate(zip(hdr_row.cells, headers)):
        cell.width = Twips(col_w[ci])
        _word_cell(cell, hdr, bold=True, font_size_pt=9,
                   color_hex="FFFFFF", align="center",
                   bg_hex="1A3A6C", border_color="1A3A6C")

    for item in results:
        category  = item.get("category", "기타 핵심 조항")
        fg, bg    = _KT_COLORS_WORD.get(category, _KT_DEFAULT_WORD)

        row = tbl.add_row()
        for ci, w in enumerate(col_w):
            row.cells[ci].width = Twips(w)

        vals   = [category, item.get("item",""), item.get("clause_ref",""),
                  item.get("original",""), item.get("korean","")]
        bolds  = [True, True, True, False, False]
        fonts  = ["맑은 고딕","맑은 고딕","Calibri","Calibri","맑은 고딕"]
        fcolors= [fg, "0F1F3D", fg, "334155", "0F1F3D"]
        bgs    = [bg, bg, bg, bg, bg]
        aligns = ["center","left","center","left","left"]

        for ci, (cell, v, b, fn, fc, bgc, al) in enumerate(
                zip(row.cells, vals, bolds, fonts, fcolors, bgs, aligns)):
            cell.width = Twips(col_w[ci])
            _word_cell(cell, v, bold=b, font_name=fn, font_size_pt=9,
                       color_hex=fc, align=al, bg_hex=bgc)

    doc.save(str(output_path))
    print(f"\n  ✓ Word 저장 완료: {output_path}")
    print(f"    - 총 {len(results)}개 항목")


# ── 메인 ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="영문 계약서 → 조항별 한국어 번역 → Excel 저장",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("input", nargs="?", help="계약서 파일 경로 (PDF 또는 DOCX)")
    parser.add_argument("-o", "--output", help="출력 Excel 파일 경로")
    parser.add_argument("--mode", choices=["full", "check"], default="full",
                        help="full=전체번역, check=핵심조항 빠른 확인")
    args = parser.parse_args()

    mode = args.mode

    print("\n" + "="*60)
    if mode == "check":
        print("  핵심 조항 빠른 확인  |  Hyundai Pharm BD")
    else:
        print("  영문 계약서 전체 번역  |  Hyundai Pharm BD")
    print("="*60)

    if args.input:
        input_path = Path(args.input)
    else:
        print("\n계약서 파일 경로를 입력하세요:")
        raw = input("  파일 경로: ").strip().strip('"').strip("'")
        input_path = Path(raw)

    if not input_path.exists():
        print(f"\n[오류] 파일을 찾을 수 없습니다: {input_path}")
        sys.exit(1)

    if args.output:
        output_path = Path(args.output)
    else:
        date_str = datetime.now().strftime("%Y%m%d")
        suffix = "_핵심조항" if mode == "check" else "_번역"
        output_path = input_path.parent / f"{input_path.stem}{suffix}_{date_str}.xlsx"

    print(f"\n[1/4] 파일 확인")
    print(f"  입력: {input_path.resolve()}")
    print(f"  출력: {output_path.resolve()}")
    print(f"  크기: {input_path.stat().st_size / 1024:.1f} KB")

    print(f"\n[2/4] 텍스트 추출")
    try:
        text = extract_text(input_path)
    except Exception as e:
        print(f"\n[오류] 텍스트 추출 실패: {e}")
        sys.exit(1)

    is_image_pdf = (input_path.suffix.lower() == ".pdf" and len(text.strip()) < 300)

    if not is_image_pdf and len(text.strip()) < 100:
        print("[오류] 추출된 텍스트가 너무 짧습니다.")
        sys.exit(1)

    print(f"\n[3/4] API 연결")
    api_key = get_api_key()
    client = anthropic.Anthropic(api_key=api_key)

    if is_image_pdf:
        print(f"\n[2.5/4] 이미지 PDF OCR 처리")
        try:
            text = extract_pdf_ocr(input_path, client)
        except Exception as e:
            print(f"\n[오류] OCR 처리 실패: {e}")
            sys.exit(1)
        if len(text.strip()) < 100:
            print("[오류] OCR 후에도 텍스트를 추출하지 못했습니다.")
            sys.exit(1)

    try:
        if mode == "check":
            # ── 핵심 조항 빠른 확인 ──
            print(f"\n[4/4] 핵심 조항 분석 (계약기간·해지 관련)")
            results = analyze_key_terms(client, text)
            if not results:
                print("\n[오류] 핵심 조항을 추출하지 못했습니다.")
                sys.exit(1)
            save_key_terms_excel(results, output_path, input_path.name)
        else:
            # ── 전체 번역 ──
            print(f"\n[4/4] Claude API 조항 추출 및 전체 번역")
            results = translate_contract(client, text)
            if not results:
                print("\n[오류] 조항을 추출하지 못했습니다.")
                sys.exit(1)
            save_excel(results, output_path, input_path.name)

    except anthropic.AuthenticationError:
        print("\n[오류] API 키가 유효하지 않습니다.")
        sys.exit(1)
    except anthropic.RateLimitError:
        print("\n[오류] API 요청 한도 초과. 잠시 후 다시 시도하세요.")
        sys.exit(1)
    except Exception as e:
        print(f"\n[오류] API 호출 실패: {e}")
        sys.exit(1)

    print("\n" + "="*60)
    print("  처리 완료!")
    print(f"  결과 파일: {output_path.resolve()}")
    print("="*60 + "\n")


if __name__ == "__main__":
    main()
